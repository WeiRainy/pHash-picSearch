#!/usr/bin/python
# -*- coding: UTF-8 -*-
# Less than 10 add to list and sort
import glob
import os
import sys
from functools import reduce
import docx
from PIL import Image

EXT = 'jpg', 'jpeg', 'JPG', 'JPEG', 'gif', 'GIF', 'png'


def avhash(im):
    if not isinstance(im, Image.Image):
        im = Image.open(im)
    if im.mode == 'RGBA':
        im = ConvertRBGA(im)
    im = im.resize((8, 8), Image.ANTIALIAS).convert('L')
    avg = reduce(lambda x, y: x + y, im.getdata()) / 64.0
    return reduce(lambda x, y_z: x | (y_z[1] << y_z[0]), enumerate(map(lambda i: 0 if i < avg else 1, im.getdata())), 0)


def hamming(h1, h2):
    h, d = 0, h1 ^ h2
    while d:
        h += 1
        d &= d - 1
    return h


def ConvertRBGA(img):
    x, y = img.size
    #   # (alpha band as paste mask).
    p = Image.new('RGBA', img.size, (255, 255, 255))
    p.paste(img, (0, 0, x, y), img)
    return p


def calc(image, search):
    h = avhash(image)
    os.chdir(search)
    images = []
    for ext in EXT:
        images.extend(glob.glob('*.%s' % ext))

    seq = []
    prog = int(len(images) > 50 and sys.stdout.isatty())
    for f in images:
        result = avhash(f)
        seq.append((f, hamming(result, h)))
        if prog:
            perc = 100. * prog / len(images)
            x = int(2 * perc / 5)
            print('\rCalculating... [' + '#' * x + ' ' * (40 - x) + ']'),
            print('%.2f%%' % perc, '(%d/%d)' % (prog, len(images))),
            sys.stdout.flush()
            prog += 1

    if prog: print()
    print("计算中......")
    print("Hamming\tFile")
    for f, ham in sorted(seq, key=lambda i: i[1]):
        print("%d\t\t%s" % (ham, f))
        if ham < 3:
            print("已找到, 文件名:" + f)
            return f.split('.')[-2]


def result(pic_path, file_path, wd):
    doc = docx.Document(file_path)
    pic_name = calc(pic_path, wd)
    print(pic_name)
    pic_order = eval(pic_name.replace('image', ''))
    para_no = 0
    pics = 0
    for para in doc.paragraphs:
        para_no += 1
        if para._element.xpath('.//pic:pic') != []:
            pics += 1
            if pics == pic_order:
                break
    hint = ""
    hint += doc.paragraphs[para_no-2].text
    hint += "\n[该位置为图片]\n"
    hint += doc.paragraphs[para_no].text
    return hint